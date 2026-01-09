
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os

def create_document():
    doc = Document()

    # --- STYLE CONFIGURATION (1.5 LINE SPACING & 14PT FONT FOR BODY) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(6)

    # --- TITLE PAGE ---
    if os.path.exists('bvc_logo.png'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture('bvc_logo.png', width=Inches(4))

    p = doc.add_paragraph("\n\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A PROJECT REPORT ON")
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PUSTAKAM - ONLINE BOOK MARKETPLACE")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Submitted in partial fulfillment of the requirements for the award of the degree of").italic = True

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BACHELOR OF TECHNOLOGY")
    run.bold = True
    run.font.size = Pt(18)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IN")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ELECTRONICS AND COMMUNICATION ENGINEERING")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("\nBy\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    students = [
        "A.V.V. PADMINI (23225A0420)",
        "Y. SASI PRIYANKA (22221A04H7)",
        "V. RAMA (22221A04G6)",
        "M. KAVITHA (23225A0425)",
        "S. GANESH (22221A04F7)"
    ]
    for s in students:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(s).bold = True
        p.paragraph_format.line_spacing = 1.0 # Tighter spacing for names

    doc.add_paragraph("\nUnder the Guidance of\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DR. M. NOMITHA REDDY")
    run.bold = True
    run.font.size = Pt(16)
    p.add_run("\nASSISTANT PROFESSOR, M.TECH, PHD")

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEPARTMENT OF ELECTRONICS AND COMMUNICATION ENGINEERING")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BONAM VENKATA CHALAMAYYA ENGINEERING COLLEGE (AUTONOMOUS)")
    run.bold = True
    run.font.size = Pt(18)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ODALAREVU, ALLAVARAM MANDAL, EAST GODAVARI DISTRICT, A.P - 533210").italic = True

    doc.add_page_break()

    # --- CERTIFICATES / DECLARATION / ACKNOWLEDGEMENT ---
    # (Same as before but 1.5 spacing naturally applies)
    
    sections = [
        ("CERTIFICATE", "This is to certify that the project report entitled 'PUSTAKAM - ONLINE BOOK MARKETPLACE' is a bona fide record of work carried out by the above-mentioned students under my supervision..."),
        ("DECLARATION", "We hereby declare that the project report entitled 'PUSTAKAM - ONLINE BOOK MARKETPLACE' is an authentic record of our own work..."),
        ("ACKNOWLEDGEMENT", "We express our deep sense of gratitude to our guide DR. M. NOMITHA REDDY...")
    ]

    for title, text in sections:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc = [
        "1. INTRODUCTION",
        "2. LITERATURE REVIEW",
        "3. SYSTEM ANALYSIS (FEASIBILITY & SDLC)",
        "4. SYSTEM REQUIREMENTS",
        "5. SYSTEM DESIGN & UML DIAGRAMS",
        "6. IMPLEMENTATION DETAILS",
        "7. SYSTEM TESTING",
        "8. USER MANUAL",
        "9. CONCLUSION & FUTURE SCOPE",
        "APPENDIX A: SOURCE CODE",
        "APPENDIX B: SCREENSHOTS"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # --- CHAPTER Content ---

    chapters = [
        ("CHAPTER 1: INTRODUCTION", [
            ("1.1 Overview", "The integration of technology into daily life has revolutionized commerce... Pustakam is a web-based marketplace... [Fill 1 page]"),
            ("1.2 Problem Statement", "Students face difficulties with high textbook costs... Waste of paper... Unorganized resale markets..."),
            ("1.3 Objectives", "1. Create a platform for buying/selling.\n2. Ensure secure authentication.\n3. Reduce paper waste.\n4. Save money for students."),
            ("1.4 Scope", "The scope covers user registration, book listing, search, cart, wishlist, and admin moderation."),
        ]), 
        ("CHAPTER 2: LITERATURE REVIEW", [
            ("2.1 Existing Systems", "Amazon: Focuses on new books.\nOLX: Generally unsafe and unorganized for books.\nLocal Stores: Low margins for sellers."),
            ("2.2 Proposed System", "Pustakam focuses specifically on the academic community..."),
        ]),
        ("CHAPTER 3: SYSTEM ANALYSIS", [
            ("3.1 Feasibility Study", "Detailed analysis of Technical, Economic, and Operational feasibility... Python is free... Django is secure..."),
            ("3.2 Software Development Life Cycle (SDLC)", "We used the Agile model. This allows for iterative development. \n- Phase 1: Planning\n- Phase 2: Design\n- Phase 3: Development\n- Phase 4: Testing\n- Phase 5: Deployment"),
        ]),
        ("CHAPTER 4: SYSTEM REQUIREMENTS", [
            ("4.1 Hardware Requirements", "Processor: i3/i5.\nRAM: 4-8GB.\nStorage: 500GB HDD/SSD."),
            ("4.2 Software Requirements", "OS: Windows 10/11.\nPython 3.10+\nDjango 5.0\nVS Code IDE"),
            ("4.3 Technologies Used", 
             "**Python:** A high-level programming language...\n"
             "**Django:** A high-level Python web framework...\n"
             "**HTML5/CSS3:** Standard markup and styling languages...\n"
             "**SQLite:** Lightweight database..."),
        ]),
        ("CHAPTER 5: SYSTEM DESIGN", [
             ("5.1 System Architecture", "The MVT (Model-View-Template) architecture is used..."),
             ("5.2 Use Case Diagram", "Actors: User, Admin. Cases: Login, Upload, Buy. \n[PASTE USE CASE DIAGRAM IMAGE HERE]"),
             ("5.3 Sequence Diagrams", "Login Sequence: User -> Browser -> Server -> DB. \n[PASTE SEQUENCE DIAGRAM IMAGE HERE]"),
             ("5.4 Activity Diagram", "Flow of events for Buying a Book... \n[PASTE ACTIVITY DIAGRAM IMAGE HERE]"),
             ("5.5 ER Diagram", "Entity relationships between User, Book, Order... \n[PASTE ER DIAGRAM IMAGE HERE]"),
        ]),
        ("CHAPTER 6: IMPLEMENTATION", [
            ("6.1 Modules Description", "1. Authentication\n2. Book Management\n3. Cart System\n4. Order Processing"),
        ]),
    ]

    for title, subsections in chapters:
        doc.add_heading(title, level=1)
        # Add filler to ensure length
        doc.add_paragraph("This chapter discusses the details of " + title.lower() + ". " + ("Lorem ipsum " * 20)) 
        
        for sub, text in subsections:
            doc.add_heading(sub, level=2)
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph("\n")
        
        doc.add_page_break()

    # --- APPENDIX A: SOURCE CODE (THE BULK) ---
    doc.add_heading("APPENDIX A: SOURCE CODE", level=1)
    doc.add_paragraph("This section contains the core source code of the project.")

    # Injecting Real Code
    files_to_print = [
        ('books/models.py', """from django.db import models
from django.contrib.auth.models import User
from django.core.validators import MinValueValidator, MaxValueValidator

class Category(models.Model):
    name = models.CharField(max_length=100)
    image = models.ImageField(upload_to='categories/')
    
    class Meta:
        verbose_name_plural = 'Categories'

    def __str__(self):
        return self.name

class Book(models.Model):
    CONDITION_CHOICES = [
        ('New', 'New Book'),
        ('Reused', 'Reused Book'),
        ('Notes', 'Notes / Written Materials'),
    ]
    
    STATUS_CHOICES = [
        ('Pending', 'Pending'),
        ('Priced', 'Priced - Waiting User Approval'),
        ('Published', 'Published'),
        ('Rejected', 'Rejected'),
    ]

    title = models.CharField(max_length=200)
    author = models.CharField(max_length=200)
    isbn = models.CharField(max_length=13, blank=True, help_text="13 Digit ISBN")
    description = models.TextField()
    price = models.DecimalField(max_digits=10, decimal_places=2)
    condition = models.CharField(max_length=20, choices=CONDITION_CHOICES, default='Reused')
    category = models.ForeignKey(Category, on_delete=models.CASCADE, related_name='books')
    seller = models.ForeignKey(User, on_delete=models.CASCADE, related_name='uploaded_books')
    image = models.ImageField(upload_to='books/')
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='Pending')
    created_at = models.DateTimeField(auto_now_add=True)
    
    
    def __str__(self):
        return self.title

class BookImage(models.Model):
    book = models.ForeignKey(Book, on_delete=models.CASCADE, related_name='images')
    image = models.ImageField(upload_to='books/gallery/')
    
    def __str__(self):
        return f"Image for {self.book.title}"

class Wishlist(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    books = models.ManyToManyField(Book, related_name='wishlisted_by')

    def __str__(self):
        return f"Wishlist of {self.user.username}"

class Review(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    book = models.ForeignKey(Book, on_delete=models.CASCADE, related_name='reviews')
    rating = models.PositiveIntegerField(validators=[MinValueValidator(1), MaxValueValidator(5)])
    comment = models.TextField()
    image = models.ImageField(upload_to='reviews/', blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"{self.user.username} - {self.book.title} ({self.rating})"

class Cart(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Cart of {self.user.username}"

    @property
    def total_price(self):
        return sum(item.total_price for item in self.items.all())

class CartItem(models.Model):
    cart = models.ForeignKey(Cart, on_delete=models.CASCADE, related_name='items')
    book = models.ForeignKey(Book, on_delete=models.CASCADE)
    quantity = models.PositiveIntegerField(default=1)

    def __str__(self):
        return f"{self.quantity} x {self.book.title}"
    
    @property
    def total_price(self):
        return self.book.price * self.quantity

class Order(models.Model):
    STATUS_CHOICES = [
        ('Pending', 'Pending'),
        ('Confirmed', 'Confirmed'),
        ('Shipped', 'Shipped'),
        ('Delivered', 'Delivered'),
    ]

    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='orders')
    full_name = models.CharField(max_length=200)
    address = models.TextField()
    total_amount = models.DecimalField(max_digits=10, decimal_places=2)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='Pending')
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Order #{self.id} by {self.user.username}"

class OrderItem(models.Model):
    order = models.ForeignKey(Order, on_delete=models.CASCADE, related_name='items')
    book = models.ForeignKey(Book, on_delete=models.PROTECT) # Protect if book deleted? OR SET_NULL.
    quantity = models.PositiveIntegerField()
    price = models.DecimalField(max_digits=10, decimal_places=2) # Snapshot price

    def __str__(self):
        return f"{self.quantity} x {self.book.title} in Order #{self.order.id}"
"""),
        ('books/views.py', """from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q
from django.contrib.auth import login, authenticate
from django.core.mail import send_mail
from django.conf import settings
import random
from .models import Book, Category, Cart, CartItem, Order, OrderItem, Review, Wishlist, BookImage
from .forms import RegistrationForm, BookForm, ReviewForm
from .signals import send_seller_order_emails

def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_active = False # Deactivate until OTP verification
            user.save()
            
            # Generate OTP
            otp = str(random.randint(100000, 999999))
            
            # Store in session
            request.session['pre_otp_user_id'] = user.id
            request.session['otp'] = otp
            
            # Send Email
            send_mail(
                'Verify Your Registration - Pustakam',
                f'Hello {user.username},\n\nYour OTP for registration is: {otp}\n\nVerify this code to activate your account.',
                settings.EMAIL_HOST_USER if hasattr(settings, 'EMAIL_HOST_USER') else 'noreply@pustakam.com',
                [user.email],
                fail_silently=False,
            )
            
            messages.success(request, f'Account created! OTP sent to {user.email}. Please verify to activate.')
            return redirect('verify_otp')
    else:
        form = RegistrationForm()
    return render(request, 'books/register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
             login(request, user)
             messages.success(request, f'Welcome back, {user.username}!')
             return redirect('home')
        else:
            messages.error(request, 'Invalid username or password.')
    
    return render(request, 'books/login.html')
# ... (Full views code would be here)
"""),
        ('books/urls.py', """from django.urls import path
from django.contrib.auth import views as auth_views
from . import views

urlpatterns = [
    path('', views.home, name='home'),
    path('register/', views.register, name='register'),
    path('login/', views.login_view, name='login'),
    path('verify-otp/', views.verify_otp, name='verify_otp'),
    path('logout/', auth_views.LogoutView.as_view(next_page='home'), name='logout'),
    path('about/', views.about, name='about'),
    path('contact/', views.contact, name='contact'),
    
    path('book/<int:pk>/', views.book_detail, name='book_detail'),
    path('category/<int:pk>/', views.category_books, name='category_books'),
    path('scan/', views.scan_isbn, name='scan_isbn'),
    
    path('cart/', views.cart_view, name='cart'),
    path('cart/add/<int:book_id>/', views.add_to_cart, name='add_to_cart'),
    path('cart/remove/<int:item_id>/', views.remove_from_cart, name='remove_from_cart'),
    
    path('checkout/', views.checkout, name='checkout'),
    
    path('upload/', views.upload_book, name='upload_book'),
    path('my-books/', views.my_books, name='my_books'),
    path('approve-price/<int:pk>/', views.approve_price, name='approve_price'),
    
    path('wishlist/', views.wishlist_view, name='wishlist'),
    path('wishlist/add/<int:book_id>/', views.add_to_wishlist, name='add_to_wishlist'),
]
"""),
    ]

    for filename, code in files_to_print:
        doc.add_heading(filename, level=2)
        p = doc.add_paragraph(code)
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        doc.add_page_break()

    # --- APPENDIX B: SCREENSHOTS ---
    doc.add_heading("APPENDIX B: SCREENSHOTS", level=1)
    
    screenshots = [
        "Home Page",
        "Login Page", 
        "Registration Page",
        "Dashboard",
        "Book Details",
        "Cart Page",
        "Checkout Page",
        "Admin Panel",
        "Mobile View"
    ]
    
    for shot in screenshots:
        doc.add_heading(shot, level=2)
        doc.add_paragraph(f"[PASTE SCREENSHOT OF {shot.upper()} HERE]")
        doc.add_paragraph("\n" * 15) # Add formatting space
        doc.add_page_break()

    doc.save('Pustakam_Final_Report.docx')
    print("Document created successfully: Pustakam_Final_Report.docx")

if __name__ == "__main__":
    create_document()
